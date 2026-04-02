from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_word_architecture():
    doc = Document()

    # Title
    title = doc.add_heading('Fall Detection System: Architecture', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Architecture Overview
    doc.add_heading('1. Overview', level=1)
    p = doc.add_paragraph('The system architecture is designed as a multi-layered real-time processing pipeline. It follows a hybrid detection approach, combining geometric computer vision with machine learning classification.')

    # 4-Layer Breakdown
    doc.add_heading('2. The 4-Layer Architecture', level=1)
    
    layers = [
        ('Input Layer', 'OpenCV captures live video frames from the camera feed.'),
        ('Perception Layer', 'MediaPipe BlazePose extracts 33 body landmarks (x, y, z coordinates) from each frame.'),
        ('Processing Layer', 'Hybrid Detection Engine: Calculates Vertical Velocity and Body Angle (Geometric Logic) and validates using a Random Forest Classifier.'),
        ('Output Layer', 'Flask Web Dashboard displays processed MJPEG stream and triggers Audio Alerts via Winsound.')
    ]

    for layer_name, description in layers:
        run = doc.add_paragraph(style='List Bullet').add_run(f'{layer_name}: ')
        run.bold = True
        doc.add_paragraph(description, style='Normal').paragraph_format.left_indent = Pt(18)

    # Data Flow Description (since we can't easily put Mermaid in Word)
    doc.add_heading('3. Data Flow', level=1)
    flow_steps = [
        'Camera captures raw video frames.',
        'OpenCV converts frames for MediaPipe processing.',
        'MediaPipe extracts skeletal landmarks.',
        'Logic Engine calculates movement velocity and pose orientation.',
        'Random Forest Model performs activity classification.',
        'Flask streams result to the Dashboard UI.',
        'System triggers Alarm if "Fall" is confirmed.'
    ]
    
    for step in flow_steps:
        doc.add_paragraph(step, style='List Number')

    # Save to Desktop
    output_path = r"C:\Users\surya\Desktop\System_Architecture.docx"
    doc.save(output_path)
    print(f"Word file saved to: {output_path}")

if __name__ == "__main__":
    create_word_architecture()
